
#!/usr/bin/env python3
"""
Script para generar el Informe TP - Aprendizaje Automático (Entrega Inicial)
mejorado, en formato .docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Pinta el fondo de una celda con el color hexadecimal indicado."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in kwargs.items():
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color_hex="1a237e"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
    return p

def add_para(doc, text, bold=False, italic=False, size=11, space_before=0, space_after=6, color=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))
    return p

def add_mixed_para(doc, parts, size=11, space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """parts = list of (text, bold, italic, color_hex_or_None)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic, color in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))
    return p

def add_bullet(doc, text, bold_prefix=None, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
        run2 = p.add_run(text)
        run2.font.size = Pt(size)
    else:
        run = p.add_run(text)
        run.font.size = Pt(size)
    return p

def add_styled_table(doc, headers, rows, col_widths=None, header_bg="1a237e", alt_bg="e8eaf6"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cells[i], header_bg)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        bg = alt_bg if ri % 2 == 0 else "ffffff"
        for ci, val in enumerate(row):
            row_cells[ci].text = str(val)
            row_cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_bg(row_cells[ci], bg)
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])
    return table

def add_note_box(doc, text, icon="📝", bg="e3f2fd", border_color="1565c0"):
    """Caja de nota/observación con fondo de color."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_border(cell, top=border_color, bottom=border_color, left=border_color, right=border_color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{icon}  {text}")
    run.font.size = Pt(10)
    run.italic = True
    doc.add_paragraph()  # espacio después
    return table


# ── DOCUMENTO ────────────────────────────────────────────────────────────────

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

# Estilo de fuente base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ─── PORTADA ─────────────────────────────────────────────────────────────────

# Título principal
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Trabajo Práctico Grupal")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(26, 35, 126)  # azul oscuro

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run("Aprendizaje Automático — Entrega Inicial")
r2.bold = True
r2.font.size = Pt(15)
r2.font.color.rgb = RGBColor(69, 90, 100)

# Línea decorativa (tabla de 1 celda)
hr = doc.add_table(rows=1, cols=1)
hr.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_bg(hr.rows[0].cells[0], "1a237e")
hr.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(3)
hr.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(3)
for section in doc.sections:
    width = section.page_width - section.left_margin - section.right_margin
hr.columns[0].width = width
doc.add_paragraph()

# Metadatos
meta = [
    ("Institución:", "Universidad Nacional de Rosario (UNR)"),
    ("Facultad:", "Facultad de Ciencias Exactas, Ingeniería y Agrimensura (FCEIA)"),
    ("Materia:", "Aprendizaje Automático"),
    ("Opción:", "Opción 1 — Dataset real con modelo de regresión supervisada"),
    ("Fecha de entrega:", "21 de mayo de 2026"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r_label = p.add_run(f"{label} ")
    r_label.bold = True
    r_label.font.size = Pt(11)
    r_value = p.add_run(value)
    r_value.font.size = Pt(11)

doc.add_paragraph()

# ─── SECCIÓN 1: CONTEXTO Y FORMULACIÓN DEL PROBLEMA ─────────────────────────
doc.add_page_break()

add_heading(doc, "1. Contexto y Formulación del Problema", level=1)

add_para(doc,
    "El presente trabajo se enmarca en la Opción 1 de la consigna: disponemos de un dataset real "
    "obtenido de una fuente de datos abiertos del sector público argentino. El problema planteado "
    "consiste en predecir el precio internacional mensual del Oro (medido en USD por onza troy) "
    "para el período siguiente, transformando la serie temporal en un problema de regresión supervisada "
    "mediante la técnica de variables de rezago (lags).",
    space_after=6
)

add_para(doc,
    "La elección del oro como activo a predecir responde a tres fundamentos concretos:",
    space_after=4
)

add_bullet(doc, "Es un activo de refugio de relevancia económica global, cuya predicción tiene "
           "aplicaciones directas para inversores, analistas financieros y tomadores de decisiones de política monetaria.")
add_bullet(doc, "Cuenta con una serie histórica continua de 36 años (enero 1990 – abril 2026, 436 observaciones), "
           "un volumen favorable para el entrenamiento de modelos de Aprendizaje Automático.")
add_bullet(doc, "Presenta correlaciones estadísticamente significativas con otros minerales presentes en el dataset, "
           "lo que permite enriquecer el espacio de features más allá de los propios rezagos del oro.")

doc.add_paragraph()

add_para(doc,
    "Tal como indicó la cátedra, Random Forest Regressor, SVR y XGBoost son modelos de regresión general "
    "y no específicos para series temporales. No obstante, pueden comportarse de manera adecuada cuando "
    "se realiza un preprocesamiento cuidadoso orientado a capturar la dependencia temporal — en particular, "
    "la construcción de variables de rezago —. Siguiendo el feedback recibido, el trabajo se centrará en "
    "la selección, entrenamiento y análisis exhaustivo de un único modelo, profundizando en el "
    "feature engineering temporal y la justificación de los resultados obtenidos.",
    space_after=8
)

# ─── SECCIÓN 2: DESCRIPCIÓN DEL DATASET ─────────────────────────────────────
add_heading(doc, "2. Descripción del Dataset", level=1)

add_heading(doc, "2.1 Fuente y Características Generales", level=2, color_hex="283593")

add_para(doc,
    "El dataset original fue obtenido del portal de Datos Abiertos del Ministerio de Economía de la "
    "Nación Argentina (MECON), bajo el título \"Precios internacionales de minerales\". Se trata de "
    "un registro de cotizaciones mensuales de 12 commodities minerales desde enero de 1990 hasta "
    "abril de 2026.",
    space_after=6
)

add_mixed_para(doc, [
    ("Enlace de acceso: ", True, False, None),
    ("https://datos.gob.ar/dataset/sspm-precios-internacionales-minerales", False, True, "1565c0"),
], space_after=8)

add_heading(doc, "2.2 Estructura del Dataset Original", level=2, color_hex="283593")

add_para(doc,
    "En su forma original, el dataset tiene estructura larga (long format): cada fila representa la "
    "cotización de un mineral en un mes determinado. La tabla a continuación describe cada columna:",
    space_after=6
)

cols_headers = ["Columna Original", "Tipo de Dato", "Descripción"]
cols_rows = [
    ["Año", "int64", "Año de la observación (rango: 1990 – 2026)."],
    ["Mes", "int64", "Mes de la observación (rango: 1 – 12)."],
    ["Fecha", "object (string)", "Identificador temporal en formato YYYYMMM (ej. 2023M05)."],
    ["Precio", "float64", "Valor internacional mensual del mineral (en la unidad indicada)."],
    ["Mineral", "object (string)", "Nombre del commodity evaluado (12 categorías distintas)."],
    ["Unidad de medida", "object (string)", "Unidad monetaria y de masa (ej. USD/troy oz, USDT/TM)."],
    ["Numero Índice", "float64", "Índice de precios normalizado respecto a un período base."],
]
add_styled_table(doc, cols_headers, cols_rows, col_widths=[4.0, 3.5, 9.0])
doc.add_paragraph()

add_para(doc,
    "El dataset contiene un total de 4.886 observaciones y 7 columnas, correspondientes a 12 minerales "
    "con series de 436 meses cada una (a excepción del Carbonato de Litio, disponible solo desde 2018).",
    space_after=8
)

add_heading(doc, "2.3 Variables Seleccionadas para el Modelado", level=2, color_hex="283593")

add_para(doc,
    "A partir del dataset original, se filtraron y trabajaron únicamente los tres minerales pertinentes "
    "al problema. La selección fue guiada por dos análisis estadísticos complementarios:",
    space_after=4
)

add_bullet(doc, "Correlación de Pearson:", bold_prefix="(i) ")
add_para(doc,
    "Mide la asociación lineal entre el precio de cada mineral y el precio del oro. "
    "Una correlación alta indica que ambas series se mueven conjuntamente en el tiempo.",
    space_after=4
)

add_bullet(doc, "Test de Causalidad de Granger:", bold_prefix="(ii) ")
add_para(doc,
    "Verifica si los valores pasados de un mineral contienen información estadísticamente útil "
    "para predecir el precio futuro del oro, más allá de lo que el propio oro predice por sí solo. "
    "Es la prueba clave para justificar la inclusión de una variable exógena en el modelo.",
    space_after=8
)

sel_headers = ["Mineral", "Correlación de Pearson (con Oro)", "Causalidad de Granger", "Decisión"]
sel_rows = [
    ["Plata (Ag)", "0.94 (muy alta)", "Significativa en todos los lags evaluados (p < 0.001)", "✅ Incluida"],
    ["Platino (Pt)", "Alta", "Significativa en lags 1, 2 y 3 (p < 0.05)", "✅ Incluida"],
    ["Níquel", "Moderada", "No significativa de forma consistente", "❌ Descartado"],
    ["Cobre (Cu)", "Moderada", "No significativa de forma consistente", "❌ Descartado"],
]
add_styled_table(doc, sel_headers, sel_rows, col_widths=[3.5, 4.5, 6.5, 2.5])
doc.add_paragraph()

add_heading(doc, "2.4 Semántica de las Variables del Modelo", level=2, color_hex="283593")

add_para(doc,
    "Una vez definido el enfoque de series temporales supervisadas mediante lags, las variables del "
    "modelo quedan definidas de la siguiente manera:",
    space_after=6
)

add_mixed_para(doc, [("Variable Objetivo (Target):", True, False, "1a237e"), (" Precio del Oro (Au) en USD/troy oz del mes ", False, False, None), ("t+1", False, True, None), (" (predicción del mes siguiente).", False, False, None)])
doc.add_paragraph()

add_para(doc, "Variables Predictoras (Features):", bold=True, color="1a237e", space_after=3)

add_bullet(doc, "Rezagos del precio del Oro (Au): precio en los meses t-1, t-2, ..., t-k.", bold_prefix="Lags propios — ")
add_bullet(doc, "Rezagos del precio de la Plata (Ag): meses previos relevantes según el análisis de Granger.", bold_prefix="Lags de Plata — ")
add_bullet(doc, "Rezagos del precio del Platino (Pt): meses t-1, t-2 y t-3, donde se verificó significancia estadística.", bold_prefix="Lags de Platino — ")

doc.add_paragraph()
add_note_box(doc,
    "Nota metodológica: La cantidad exacta de lags a incluir (ventana óptima k) se determinará "
    "en la siguiente etapa mediante análisis de autocorrelación (ACF/PACF) y validación cruzada "
    "temporal (Time Series Split), evitando el data leakage.",
    icon="📌", bg="fff8e1", border_color="f57f17"
)

# ─── SECCIÓN 3: TAREAS DE LIMPIEZA Y ANÁLISIS EXPLORATORIO ───────────────────
add_heading(doc, "3. Tareas de Limpieza y Análisis Exploratorio", level=1)

add_para(doc,
    "Previo a la estructuración como problema supervisado, se ejecutaron las siguientes tareas de "
    "auditoría y limpieza sobre el dataset original mediante una libreta de Jupyter Notebook "
    "(TP_Aprendizaje_Automático.ipynb):",
    space_after=8
)

# 3.1 Nulos y duplicados
add_heading(doc, "3.1 Verificación de Nulos y Duplicados", level=2, color_hex="283593")

add_para(doc,
    "Se auditó la matriz completa (4.886 observaciones × 7 columnas). El resultado fue:",
    space_after=4
)
add_bullet(doc, "Valores nulos (missing values): 0 detectados en todas las columnas.", bold_prefix="✅ ")
add_bullet(doc, "Filas duplicadas: 0 detectadas.", bold_prefix="✅ ")
add_bullet(doc, "Tipo de dato de cada columna: consistente y acorde a su semántica.", bold_prefix="✅ ")
add_para(doc,
    "La ausencia de valores faltantes elimina la necesidad de aplicar técnicas de imputación, "
    "lo que simplifica el pipeline de preprocesamiento y evita introducir sesgos artificiales en la serie.",
    space_after=8
)

# 3.2 Continuidad temporal
add_heading(doc, "3.2 Control de Continuidad Temporal", level=2, color_hex="283593")

add_para(doc,
    "Al tratarse de un problema de predicción de series temporales, la continuidad de la secuencia "
    "mensual es un requisito crítico: la presencia de saltos temporales invalidaría los rezagos "
    "calculados como features. Se implementó un algoritmo que construye el conjunto de pares "
    "(Año, Mes) esperados en el rango del mineral y los contrasta con los pares presentes en el dataset.",
    space_after=6
)

add_para(doc, "Resultado para los tres minerales de interés:", bold=True, space_after=4)

cont_headers = ["Mineral", "Período cubierto", "Observaciones", "Resultado"]
cont_rows = [
    ["Oro (Au)", "Ene 1990 – Abr 2026", "436", "✅ Serie completa, sin saltos"],
    ["Plata (Ag)", "Ene 1990 – Abr 2026", "436", "✅ Serie completa, sin saltos"],
    ["Platino (Pt)", "Ene 1990 – Abr 2026", "436", "✅ Serie completa, sin saltos"],
]
add_styled_table(doc, cont_headers, cont_rows, col_widths=[3.5, 5.0, 3.5, 5.0])
doc.add_paragraph()

add_para(doc,
    "El año 2026 figura con solo 4 meses (enero–abril), lo cual es esperado y correcto dado que "
    "el dataset fue actualizado hasta dicha fecha. No constituye un salto temporal.",
    space_after=8
)

# 3.3 Outliers
add_heading(doc, "3.3 Análisis de Valores Atípicos (Outliers)", level=2, color_hex="283593")

add_para(doc,
    "Mediante la visualización de un diagrama de boxplot agrupado por año, se identificaron picos "
    "de volatilidad pronunciados en la cotización del oro, particularmente en:",
    space_after=4
)
add_bullet(doc, "2008-2009: crisis financiera global (Lehman Brothers).")
add_bullet(doc, "2011-2012: máximo histórico previo, impulsado por la crisis de deuda soberana europea.")
add_bullet(doc, "2020-2022: pandemia COVID-19, posterior estímulo monetario y conflicto en Ucrania.")
add_bullet(doc, "2024-2026: nuevo máximo histórico, superando los USD 5.000/troy oz.")

add_para(doc,
    "Estos valores no se imputarán ni eliminarán. Son fluctuaciones reales del mercado financiero "
    "con causas macroeconómicas identificables, no errores de medición o carga. Eliminarlos "
    "distorsionaría la distribución real del fenómeno y perjudicaría la capacidad predictiva del modelo "
    "en contextos de alta volatilidad, que son precisamente los más relevantes para los inversores.",
    space_after=8
)

# ─── SECCIÓN 4: ESTADÍSTICAS DESCRIPTIVAS ────────────────────────────────────
add_heading(doc, "4. Estadísticas Descriptivas de las Series Seleccionadas", level=1)

add_para(doc,
    "A continuación se presenta el resumen estadístico de los tres minerales que integrarán el modelo, "
    "calculado sobre la totalidad de sus 436 observaciones mensuales:",
    space_after=6
)

desc_headers = ["Mineral", "Obs.", "Media (USD)", "Desv. Est.", "Mínimo", "P25", "Mediana", "P75", "Máximo"]
desc_rows = [
    ["Oro (Au) — USD/troy oz",    "436", "1.017,15", "821,71",  "256,08",  "339,55", "686,68", "1.527,87", "5.019,97"],
    ["Plata (Ag) — USD/troy oz",  "436",    "14,68",  "11,92",    "3,65",    "4,89",   "12,55",     "19,87",    "92,06"],
    ["Platino (Pt) — USD/troy oz","436",   "896,19", "444,95",  "341,19",  "481,69",  "895,31",  "1.189,50", "2.433,83"],
]
add_styled_table(doc, desc_headers, desc_rows,
    col_widths=[5.5, 1.5, 2.5, 2.5, 2.0, 2.0, 2.2, 2.2, 2.2])
doc.add_paragraph()

add_para(doc,
    "Observaciones relevantes a partir de la estadística descriptiva:",
    bold=True, space_after=4
)
add_bullet(doc, "El precio del Oro muestra una desviación estándar cercana a su media (coeficiente de "
           "variación ≈ 80,8 %), lo que evidencia una alta dispersión histórica y una fuerte tendencia alcista "
           "en el largo plazo.")
add_bullet(doc, "La Plata presenta el mayor coeficiente de variación relativo a su escala, con un máximo "
           "(USD 92,06) que casi quintuplica su valor mediano, impulsado por episodios especulativos puntuales.")
add_bullet(doc, "El Platino, si bien comparte la unidad de medida (USD/troy oz), muestra una distribución "
           "más simétrica y menor dispersión relativa, lo que sugiere una dinámica de mercado diferente "
           "a la del oro.")

doc.add_paragraph()

# ─── SECCIÓN 5: CONCLUSIÓN ────────────────────────────────────────────────────
add_heading(doc, "5. Conclusión del Preprocesamiento y Próximos Pasos", level=1)

add_para(doc,
    "Las tareas de auditoría confirman que el dataset original posee integridad estructural completa: "
    "ausencia de valores nulos, sin duplicados, y con continuidad temporal perfecta para los tres "
    "minerales de interés (Oro, Plata y Platino) durante los 436 meses del período enero 1990 – "
    "abril 2026. Estas condiciones son óptimas para la construcción de un modelo de regresión "
    "supervisada basado en lags temporales.",
    space_after=8
)

add_para(doc, "Las etapas planificadas para la siguiente entrega son:", bold=True, space_after=4)

add_bullet(doc, "Feature Engineering: construcción del dataset supervisado mediante variables de rezago "
           "(lag_1 a lag_k) para Oro, Plata y Platino; determinación de k mediante ACF/PACF.", bold_prefix="Etapa 1 — ")
add_bullet(doc, "División temporal: separación cronológica en conjunto de entrenamiento y prueba "
           "(sin shuffle aleatorio, respetando el orden temporal para evitar data leakage).", bold_prefix="Etapa 2 — ")
add_bullet(doc, "Selección y justificación del modelo: se adoptará Random Forest Regressor como "
           "algoritmo único, justificando la elección por su robustez ante outliers, su capacidad "
           "de capturar relaciones no lineales entre rezagos y su buena performance documentada "
           "en series temporales transformadas a formato supervisado.", bold_prefix="Etapa 3 — ")
add_bullet(doc, "Evaluación: métricas de error (MAE, RMSE, R²) sobre el conjunto de prueba, "
           "análisis de residuos e interpretación del modelo (feature importance).", bold_prefix="Etapa 4 — ")

doc.add_paragraph()
add_note_box(doc,
    "Elección de modelo: Se opta por Random Forest Regressor por su robustez ante valores atípicos "
    "(característica esencial dada la alta volatilidad histórica del oro), su capacidad intrínseca de "
    "capturar interacciones no lineales entre rezagos y su resistencia al sobreajuste mediante el "
    "mecanismo de bagging. Adicionalmente, la importancia de features que provee el algoritmo "
    "permitirá interpretar qué rezagos y qué minerales resultan más predictivos.",
    icon="🌲", bg="e8f5e9", border_color="2e7d32"
)

# ─── GUARDAR ─────────────────────────────────────────────────────────────────
output_path = "/home/bn/Documents/UGR/AA/TP-Final/Informe_TP_AA_EntregaInicial_v2.docx"
doc.save(output_path)
print(f"✅ Informe guardado en: {output_path}")
